"""
Document Search Tool
========================

A robust, high-performance utility for searching text patterns across Microsoft Word and PDF
documents. This tool provides both command-line and interactive interfaces to efficiently search
through large document collections with export capabilities.

Features:
---------
* Support for multiple document formats:
  - Microsoft Word (DOCX)
  - PDF documents (text-based)
* Parallel processing for faster searches across multiple documents
* Recursive directory traversal with configurable exclusion patterns
* Comprehensive document coverage:
  - Word: paragraphs, tables, headers, footers
  - PDF: full text extraction from all pages
* Flexible search options:
  - Case-sensitive matching
  - Whole word matching
  - Regular expression support
* Interactive menu for user-friendly operation
* Optimized thread management based on system capabilities
* Progress reporting for large directory searches
* Export to multiple formats (HTML, Markdown, Text)
* Highlighted search results in both console and exports

Usage Examples:
--------------
Command line:
    python automated_word_search.py "search term" /path/to/directory [options]

    Options:
        --case-sensitive    Enable case-sensitive matching
        --whole-word        Match whole words only
        --regex             Use regular expression patterns
        --threads N         Set specific number of worker threads (1-32)
        --exclude DIR       Exclude directories from search (can be used multiple times)
        --pdf               Include PDF files in search (Word DOCX files always included)
        --export FORMAT     Export results in specified format: html, markdown, txt (default: none)

Interactive:
    python automated_word_search.py
    (Follow the prompts to configure your search)

Return Values:
------------
* Exit code 0: Search completed successfully (with or without matches)
* Prints matched file paths to standard output
* Exports results to script directory if specified

Dependencies:
------------
* python-docx >= 0.8.10 - For parsing Word documents
* pypdf >= 3.0.0 - For parsing PDF documents
* tqdm >= 4.65.0 - For progress reporting
* colorama >= 0.4.6 - For terminal color support
* Standard library modules: re, os, concurrent.futures, pathlib, argparse
"""

import datetime
import html
import os
import re
import sys
import time
import traceback
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from multiprocessing import cpu_count
from pathlib import Path
from typing import Any, List, Optional, Set, Tuple, TypedDict
from typing import Pattern as RegexPattern

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pypdf import PdfReader

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from tqdm import tqdm  # type: ignore  # Missing stubs

    TQDM_AVAILABLE = True
except ImportError:
    TQDM_AVAILABLE = False

try:
    from colorama import Fore, Style, init

    init(autoreset=True)
    COLORAMA_AVAILABLE = True
except ImportError:
    COLORAMA_AVAILABLE = False


@dataclass
class SearchMatch:
    """Data class to store information about each match."""

    file_path: Path
    context: str
    page_or_section: Optional[str] = None  # For PDF pages or Word document sections
    # Use field(default_factory=list) for mutable default values
    match_positions: List[Tuple[int, int]] = field(default_factory=list)


class ResultExporter:
    """Handles exporting search results to different formats."""

    def __init__(self, search_term: str, directory: Path):
        """
        Initialize the exporter with search parameters.

        Args:
            search_term: The search term or pattern used
            directory: The directory that was searched
        """
        self.search_term = search_term
        self.directory = directory
        self.timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        self.script_dir = Path(__file__).parent.resolve()

    @staticmethod
    def _wrap_text(text: str, max_width: int = 100) -> str:
        """
        Wrap text to the specified width without breaking words.

        Args:
            text: The text to wrap
            max_width: Maximum line width (default: 100)

        Returns:
            String with appropriate line breaks
        """
        lines = []
        for paragraph in text.split("\n"):
            if not paragraph:
                lines.append("")
                continue

            current_line: list[str] = []
            current_length = 0

            for word in paragraph.split():
                word_length = len(word)

                if current_length + word_length + (1 if current_line else 0) <= max_width:
                    # Word fits on current line
                    if current_line:
                        current_length += 1  # Space before word
                    current_line.append(word)
                    current_length += word_length
                else:
                    # Start a new line
                    lines.append(" ".join(current_line))
                    current_line = [word]
                    current_length = word_length

            if current_line:
                lines.append(" ".join(current_line))

        return "\n".join(lines)

    def _get_export_path(self, extension: str) -> Path:
        """Generate a file path for the export file."""
        safe_search_term = re.sub(r"[^\w\s-]", "", self.search_term)[:30]
        safe_search_term = re.sub(r"[-\s]+", "-", safe_search_term).strip("-")
        filename = f"search_results_{safe_search_term}_{self.timestamp}.{extension}"
        return self.script_dir / filename

    def export_html(self, results: List[SearchMatch]) -> Path:
        """
        Export results to HTML format with highlighted matches.

        Args:
            results: List of search matches

        Returns:
            Path to the exported file
        """
        output_path = self._get_export_path("html")

        with open(output_path, "w", encoding="utf-8") as f:
            # HTML header with proper line breaks to stay within 100 chars
            f.write(
                f"""<!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Search Results: {html.escape(self.search_term)}</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 20px; line-height: 1.6; }}
            h1 {{ color: #333; }}
            .search-info {{
                background-color: #f5f5f5;
                padding: 10px;
                border-radius: 5px;
                margin-bottom: 20px;
            }}
            .file-header {{
                background-color: #e0e0e0;
                padding: 10px;
                margin-top: 30px;
                border-radius: 5px;
            }}
            .match-location {{ color: #666; font-style: italic; margin-top: 15px; }}
            .match-context {{
                background-color: #f9f9f9;
                padding: 15px;
                border-left: 4px solid #ccc;
                white-space: pre-wrap;
            }}
            .highlight {{ background-color: #ffff99; font-weight: bold; }}
            .summary {{ margin-bottom: 20px; }}
            footer {{
                margin-top: 30px;
                color: #666;
                font-size: 0.9em;
                border-top: 1px solid #eee;
                padding-top: 10px;
            }}
        </style>
    </head>
    <body>
        <h1>Search Results</h1>
        <div class="search-info">
            <p><strong>Search Term:</strong> {html.escape(self.search_term)}</p>
            <p><strong>Directory:</strong> {html.escape(str(self.directory))}</p>
            <p><strong>Date:</strong> {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>
        </div>
        <div class="summary">
            <p>Found {len(results)} matches across {len({match.file_path for match in results})} 
            documents.</p>
        </div>
    """
            )

            current_file = None
            for result in results:
                if current_file != result.file_path:
                    if current_file is not None:
                        f.write("    <hr>\n")
                    file_path_escaped = html.escape(str(result.file_path))
                    f.write('    <div class="file-header">\n')
                    f.write(f"        <h2>{file_path_escaped}</h2>\n")
                    f.write("    </div>\n")
                    current_file = result.file_path

                location_text = html.escape(result.page_or_section or "")
                f.write(f'    <div class="match-location">{location_text}</div>\n')

            return output_path

    def export_markdown(self, results: List[SearchMatch]) -> Path:
        """
        Export results to Markdown format.

        Args:
            results: List of search matches

        Returns:
            Path to the exported file
        """
        output_path = self._get_export_path("md")

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(f'# Search Results: "{self.search_term}"\n\n')
            f.write(f"**Directory:** {self.directory}\n")
            f.write(f'**Date:** {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n\n')
            f.write(
                f"Found {len(results)} matches across "
                f"{len({match.file_path for match in results})} documents.\n\n"
            )

            current_file = None
            for result in results:
                if current_file != result.file_path:
                    f.write(f"\n## {result.file_path}\n\n")
                    current_file = result.file_path

                f.write(f'### {result.page_or_section or ""}\n\n')
                f.write("```\n")
                # Apply text wrapping to ensure lines don't exceed 100 characters
                f.write(self._wrap_text(result.context))
                f.write("\n```\n\n")

        return output_path

    def export_text(self, results: List[SearchMatch]) -> Path:
        """
        Export results to plain text format with proper line wrapping.

        Args:
            results: List of search matches

        Returns:
            Path to the exported file
        """
        output_path = self._get_export_path("txt")

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(f'Search Results: "{self.search_term}"\n')
            f.write("=" * 50 + "\n\n")
            f.write(f"Directory: {self.directory}\n")
            f.write(f'Date: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n\n')
            f.write(
                f"Found {len(results)} matches across "
                f"{len({match.file_path for match in results})} documents.\n\n"
            )

            current_file = None
            for result in results:
                if current_file != result.file_path:
                    f.write("\n" + "=" * 50 + "\n")
                    f.write(f"{result.file_path}\n")
                    f.write("=" * 50 + "\n\n")
                    current_file = result.file_path

                f.write(f'\n{result.page_or_section or ""}:\n')
                f.write("-" * 40 + "\n")
                # Use the text wrapping function here
                f.write(self._wrap_text(result.context) + "\n")
                f.write("-" * 40 + "\n")

        return output_path

    def export(self, results: List[SearchMatch], format_type: str) -> Optional[Path]:
        """
        Export results to the specified format.

        Args:
            results: List of search matches
            format_type: The format to export to (html, markdown, txt)

        Returns:
            Path to the exported file or None if format not supported
        """
        if not results:
            print("No results to export.")
            return None

        format_type = format_type.lower()

        if format_type == "html":
            return self.export_html(results)
        elif format_type in ("markdown", "md"):
            return self.export_markdown(results)
        elif format_type in ("text", "txt"):
            return self.export_text(results)
        else:
            print(f"Unsupported export format: {format_type}")
            return None


class WordSearcher:
    """
    Search for patterns in Word documents with configurable matching options.

    This class provides efficient pattern matching across Microsoft Word documents,
    searching through paragraphs, tables, headers, and footers. The search operations
    can be parallelized for improved performance on large document collections.

    Attributes:
        search_term: The word or regex pattern to search for in documents.
        case_sensitive: If True, the search is case-sensitive.
        whole_word: If True, match only whole words.
        use_regex: If True, interpret `search_term` as a regex pattern.
        max_workers: Maximum number of worker threads.
        pattern: The compiled regex pattern used for searching.
    """

    def __init__(
        self,
        search_term: str,
        case_sensitive: bool = False,
        whole_word: bool = False,
        use_regex: bool = False,
        max_workers: Optional[int] = None,
    ):
        """
        Initialize the WordSearcher with search configuration.

        Args:
            search_term: The word or regex pattern to search for in documents.
            case_sensitive: If True, the search is case-sensitive.
            whole_word: If True, match only whole words (adds word boundaries to pattern).
            use_regex: If True, interpret `search_term` as a regex pattern.
            max_workers: Maximum number of worker threads. Default is CPU count + 4,
                         capped at 32 for resource management.
        """
        self.search_term = search_term
        self.case_sensitive = case_sensitive
        self.whole_word = whole_word
        self.use_regex = use_regex
        # Set thread count based on system capabilities with a reasonable upper bound
        self.max_workers = max_workers or min(32, (cpu_count() or 4) + 4)

        # Compile the search pattern once for reuse
        self.pattern = self._compile_pattern()

    def _compile_pattern(self) -> RegexPattern[str]:
        """
        Compile the regex pattern for searching based on user options.

        Returns:
            A compiled regex pattern object configured according to search options.
        """
        # Escape special characters if not using regex mode
        term = self.search_term if self.use_regex else re.escape(self.search_term)

        # Add word boundary markers if whole word matching is enabled
        if self.whole_word:
            term = rf"\b{term}\b"

        # Set case sensitivity flag
        flags = 0 if self.case_sensitive else re.IGNORECASE
        return re.compile(term, flags)

    def _find_matches_in_text(self, text: str) -> List[Tuple[int, int]]:
        """
        Find all matches of the pattern in the given text.

        Args:
            text: The text content to search within.

        Returns:
            List of tuples containing (start, end) positions of matches.
        """
        return [(m.start(), m.end()) for m in self.pattern.finditer(text)]

    def _search_text(self, text: str) -> bool:
        """
        Check if the search pattern matches a given text.

        Args:
            text: The text content to search within.

        Returns:
            True if the pattern is found in the text, False otherwise.
        """
        return bool(self.pattern.search(text))

    def search_document(self, file_path: Path) -> List[SearchMatch]:
        """Search for the pattern in a Word document and return matches with context."""
        if not DOCX_AVAILABLE:
            print("python-docx module not installed. Cannot search Word documents.")
            return []

        matches = []
        try:
            doc = Document(str(file_path))

            # Search paragraphs in the main document body
            for i, p in enumerate(doc.paragraphs):
                if self._search_text(p.text):
                    match_positions = self._find_matches_in_text(p.text)
                    matches.append(
                        SearchMatch(
                            file_path=file_path,
                            context=p.text,
                            page_or_section=f"Paragraph {i + 1}",
                            match_positions=match_positions,
                        )
                    )

            # Search through all cells in all tables
            for t_idx, table in enumerate(doc.tables):
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        if self._search_text(cell.text):
                            match_positions = self._find_matches_in_text(cell.text)
                            matches.append(
                                SearchMatch(
                                    file_path=file_path,
                                    context=cell.text,
                                    page_or_section=(
                                        f"Table {t_idx + 1}, "
                                        f"Row {r_idx + 1}, "
                                        f"Column {c_idx + 1}"
                                    ),
                                    match_positions=match_positions,
                                )
                            )

            # Search through headers and footers
            for s_idx, section in enumerate(doc.sections):
                for hf_type, header_footer in [
                    ("Header", section.header),
                    ("Footer", section.footer),
                ]:
                    if header_footer:
                        for p_idx, p in enumerate(header_footer.paragraphs):
                            if self._search_text(p.text):
                                match_positions = self._find_matches_in_text(p.text)
                                matches.append(
                                    SearchMatch(
                                        file_path=file_path,
                                        context=p.text,
                                        page_or_section=(
                                            f"Section {s_idx + 1}, "
                                            f"{hf_type}, "
                                            f"Paragraph {p_idx + 1}"
                                        ),
                                        match_positions=match_positions,
                                    )
                                )

        except Exception as e:
            print(f"Error processing '{file_path}': {e}")
            if os.environ.get("DEBUG"):
                traceback.print_exc()

        return matches

    def search_pdf(self, file_path: Path) -> List[SearchMatch]:
        """Search for the pattern in a PDF document and return matches with context."""
        if not PDF_AVAILABLE:
            print("pypdf module not installed. Cannot search PDF documents.")
            return []

        matches = []
        try:
            with open(file_path, "rb") as file:
                reader = PdfReader(file)
                # Search through each page
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text and self._search_text(text):
                        # Split text into paragraphs to provide better context
                        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
                        for para in paragraphs:
                            if self._search_text(para):
                                match_positions = self._find_matches_in_text(para)
                                matches.append(
                                    SearchMatch(
                                        file_path=file_path,
                                        context=para,
                                        page_or_section=f"Page {page_num + 1}",
                                        match_positions=match_positions,
                                    )
                                )

        except Exception as e:
            print(f"Error processing PDF '{file_path}': {e}")
            if os.environ.get("DEBUG"):
                traceback.print_exc()

        return matches

    def process_file(self, file_path: Path) -> List[SearchMatch]:
        """Process a file based on its extension."""
        suffix = file_path.suffix.lower()
        if suffix == ".docx":
            return self.search_document(file_path)
        elif suffix == ".pdf":
            return self.search_pdf(file_path)
        return []

    @staticmethod
    def _collect_files(
        directory: Path, file_patterns: List[str], exclude_dirs: Set[str]
    ) -> List[Path]:
        """
        Collect all matching files in a directory tree.

        Args:
            directory: Root directory to search
            file_patterns: List of glob patterns to match files
            exclude_dirs: Set of directory names to exclude

        Returns:
            List of matching file paths
        """
        files = []
        for root, dirs, _ in os.walk(directory):
            # Exclude directories
            dirs[:] = [d for d in dirs if d not in exclude_dirs]

            # Match files in current directory
            for pattern in file_patterns:
                matched_files = list(Path(root).glob(pattern))
                files.extend(matched_files)

        return files

    def search_recursive(
        self,
        root_directory: Path,
        file_patterns: Optional[List[str]] = None,
        exclude_dirs: Optional[Set[str]] = None,
    ) -> List[SearchMatch]:
        """
        Recursively search for patterns in document files across directories.

        Args:
            root_directory: The root directory to start searching from.
            file_patterns: List of file patterns to match.
            exclude_dirs: Set of directory names to exclude from the search.

        Returns:
            List of matches with context information.
        """
        # Default parameters
        patterns = ["*.docx", "*.pdf"] if file_patterns is None else file_patterns
        exclude_set = (
            {".git", "__pycache__", "venv", ".venv"} if exclude_dirs is None else exclude_dirs
        )

        # Ensure we're working with an absolute path to avoid confusion
        root_directory = root_directory.resolve()
        print(f"Starting search in: {root_directory} (absolute path)")

        # First collect all files to process
        print("Collecting files to search...")
        files = WordSearcher._collect_files(root_directory, patterns, exclude_set)

        if not files:
            print("No matching files found.")
            return []

        print(f"Found {len(files)} files to process.")
        all_matches = []

        # Progress tracking
        if TQDM_AVAILABLE:
            files_to_process = tqdm(files, desc="Searching files", unit="file")
        else:
            print(f"Processing {len(files)} files...")
            files_to_process = files

        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            # Use executor.map to process files in parallel
            future_to_file = {
                executor.submit(self.process_file, file): file for file in files_to_process
            }
            for future in as_completed(future_to_file):
                matches = future.result()
                all_matches.extend(matches)
                if not TQDM_AVAILABLE:
                    file = future_to_file[future]
                    print(f"Processed: {file.name} - Found {len(matches)} matches")

        return all_matches


def highlight_text(text: str, positions: List[Tuple[int, int]]) -> str:
    """
    Highlight text at the specified positions.

    Args:
        text: Original text
        positions: List of (start, end) positions to highlight

    Returns:
        Text with highlighted portions
    """
    if not COLORAMA_AVAILABLE:
        return text

    # Sort positions in reverse to avoid offsetting earlier positions
    sorted_positions = sorted(positions, key=lambda x: x[0], reverse=True)
    result = text

    for start, end in sorted_positions:
        if start < 0 or end > len(text) or start >= end:
            continue

        highlighted = f"{Fore.YELLOW}{Style.BRIGHT}{text[start:end]}{Style.RESET_ALL}"
        result = result[:start] + highlighted + result[end:]

    return result


class SearchParams(TypedDict):
    """Type definition for search parameters."""

    search_term: str
    directory: Path
    case_sensitive: bool
    whole_word: bool
    use_regex: bool
    threads: Optional[int]
    exclude: Set[str]
    file_patterns: List[str]
    export_format: Optional[str]


def interactive_menu() -> dict[str, Any]:
    """
    Display an interactive menu to configure the search parameters.

    Returns:
        Dictionary of search parameters configured by user input
    """
    print("\n===== Word Document Search Tool =====")

    # Get search term - required field with validation
    search_term = input("\nEnter the word or pattern to search for: ").strip()
    while not search_term:
        search_term = input("Search term cannot be empty. Please enter a search term: ").strip()

    # Get directory with improved options
    script_dir = Path(__file__).parent.resolve()
    cwd = Path.cwd().resolve()

    print("\nDirectory options:")
    print(f"1. Script location: {script_dir}")
    print(f"2. Current working directory: {cwd}")
    print("3. Enter custom path")

    choice = input("Select directory option (1-3, default: 1): ").strip()

    # Use ternary operator instead of if/else blocks
    directory_path = (
        script_dir
        if not choice or choice == "1"
        else (
            cwd
            if choice == "2"
            else Path(input("Enter directory path: ").strip() or str(script_dir)).resolve()
        )
    )

    print(f"Search will be performed in: {directory_path}")
    while not directory_path.exists() or not directory_path.is_dir():
        directory = input("Invalid directory. Please enter a valid directory path: ").strip()
        directory_path = script_dir if not directory else Path(directory).resolve()

    # Get search options
    case_sensitive = input("\nEnable case-sensitive matching? (y/N): ").strip().lower() == "y"
    whole_word = input("Match whole words only? (y/N): ").strip().lower() == "y"
    use_regex = input("Interpret the search term as a regex? (y/N): ").strip().lower() == "y"

    # Get excluded directories
    print("\nEnter directories to exclude (leave blank to finish):")
    exclude_dirs = set()
    while True:
        exclude_dir = input("Exclude directory (or press Enter to continue): ").strip()
        if not exclude_dir:
            break
        exclude_dirs.add(exclude_dir)

    # Add default excluded directories
    exclude_dirs |= {".git", "__pycache__", "venv", ".venv"}

    # Get number of threads
    threads = None
    custom_threads = input("\nSpecify number of worker threads? (y/N): ").strip().lower() == "y"
    if custom_threads:
        while True:
            try:
                threads = int(input("Enter number of threads (1-32): ").strip())
                if 1 <= threads <= 32:
                    break
                print("Please enter a number between 1 and 32.")
            except ValueError:
                print("Please enter a valid number.")

    # Add option for PDF files
    include_pdf = input("\nInclude PDF files in search? (Y/n): ").strip().lower() != "n"
    file_patterns = ["*.docx"]
    if include_pdf:
        file_patterns.append("*.pdf")

    # Export options
    export_format = None
    if input("\nExport results to a file? (y/N): ").strip().lower() == "y":
        print("\nExport formats:")
        print("1. HTML (with highlighted matches)")
        print("2. Markdown")
        print("3. Plain Text")

        export_choice = input("Select export format (1-3, default: 1): ").strip()
        if not export_choice or export_choice == "1":
            export_format = "html"
        elif export_choice == "2":
            export_format = "markdown"
        elif export_choice == "3":
            export_format = "txt"
        else:
            export_format = "html"

    return {
        "search_term": search_term,
        "directory": directory_path,
        "case_sensitive": case_sensitive,
        "whole_word": whole_word,
        "use_regex": use_regex,
        "threads": threads,
        "exclude": exclude_dirs,
        "file_patterns": file_patterns,
        "export_format": export_format,
    }


def display_results(results: List[SearchMatch]) -> None:
    """Format and display search results in the console with highlighting."""
    if results:
        print(
            f"\nFound {len(results)} matches across "
            f"{len({match.file_path for match in results})} documents:"
        )
        current_file = None
        for result in results:
            if current_file != result.file_path:
                print(
                    f"\n{Fore.CYAN if COLORAMA_AVAILABLE else ''}=== "
                    f"{result.file_path} "
                    f"==={Style.RESET_ALL if COLORAMA_AVAILABLE else ''}"
                )
                current_file = result.file_path

            print(
                f"\n{Fore.BLUE if COLORAMA_AVAILABLE else ''}"
                f"{result.page_or_section if result.page_or_section else ''}"
                f"{Style.RESET_ALL if COLORAMA_AVAILABLE else ''}:"
            )
            print("-" * 40)
            if COLORAMA_AVAILABLE and result.match_positions:
                print(highlight_text(result.context, result.match_positions))
            else:
                print(result.context)
            print("-" * 40)
    else:
        print("\nNo matches found.")


def main() -> int:
    """
    Entry point for the Word Document Search Tool.

    Handles both command-line and interactive modes.
    After configuration, performs the search operation and displays results.

    Returns:
        Exit code (0 for successful execution)
    """
    import argparse

    # Check for dependencies
    if not DOCX_AVAILABLE:
        print(
            "Warning: python-docx module not installed. Word document search will be unavailable."
        )
    if not PDF_AVAILABLE:
        print("Warning: pypdf module not installed. PDF document search will be unavailable.")
    if not TQDM_AVAILABLE:
        print("Warning: tqdm module not installed. Progress reporting will be limited.")
    if not COLORAMA_AVAILABLE:
        print("Warning: colorama module not installed. Text highlighting will be unavailable.")

    # Determine whether to use CLI or interactive mode
    if len(sys.argv) > 1:
        # Parse command-line arguments
        parser = argparse.ArgumentParser(
            description="Search for patterns in Word documents.",
            formatter_class=argparse.ArgumentDefaultsHelpFormatter,
        )
        parser.add_argument("search_term", help="The word or pattern to search for")
        parser.add_argument("directory", help="The root directory to search in")
        parser.add_argument(
            "--case-sensitive", action="store_true", help="Enable case-sensitive matching"
        )
        parser.add_argument("--whole-word", action="store_true", help="Match whole words only")
        parser.add_argument("--regex", action="store_true", help="Use regular expression patterns")
        parser.add_argument(
            "--threads", type=int, default=None, help="Set specific number of worker threads (1-32)"
        )
        parser.add_argument(
            "--exclude",
            action="append",
            default=[],
            help="Exclude directories from search (can be used multiple times)",
        )
        parser.add_argument(
            "--pdf",
            action="store_true",
            help="Include PDF files in search (Word DOCX files always included)",
        )
        parser.add_argument(
            "--export",
            choices=["html", "markdown", "txt"],
            help="Export results in specified format",
        )

        args = parser.parse_args()

        # Configure search parameters from CLI arguments
        search_params = {
            "search_term": args.search_term,
            "directory": Path(args.directory).resolve(),
            "case_sensitive": args.case_sensitive,
            "whole_word": args.whole_word,
            "use_regex": args.regex,
            "threads": args.threads,
            "exclude": set(args.exclude) | {".git", "__pycache__", "venv", ".venv"},
            "file_patterns": ["*.docx"] + (["*.pdf"] if args.pdf else []),
            "export_format": args.export,
        }
    else:
        # Use interactive menu
        search_params = interactive_menu()

    start_time = time.time()

    # Create searcher with configured parameters
    searcher = WordSearcher(
        search_term=search_params["search_term"],
        case_sensitive=search_params["case_sensitive"],
        whole_word=search_params["whole_word"],
        use_regex=search_params["use_regex"],
        max_workers=search_params["threads"],
    )

    # Perform the search
    print("\nStarting search...")
    results = searcher.search_recursive(
        root_directory=search_params["directory"],
        file_patterns=search_params["file_patterns"],
        exclude_dirs=search_params["exclude"],
    )

    # Display results
    display_results(results)

    # Export results if requested
    if search_params["export_format"] and results:
        exporter = ResultExporter(
            search_term=search_params["search_term"], directory=search_params["directory"]
        )
        output_path = exporter.export(results, search_params["export_format"])
        if output_path:
            print(f"\nResults exported to: {output_path}")

    elapsed_time = time.time() - start_time
    print(f"\nSearch completed in {elapsed_time:.2f} seconds.")

    return 0


if __name__ == "__main__":
    sys.exit(main())
